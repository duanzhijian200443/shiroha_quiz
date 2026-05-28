#!/usr/bin/env python3
"""
Shiroha Quiz — PDF / DOCX → JSON 题库自动抽取引擎
依赖: pip install PyMuPDF python-docx openai
用法: python pdf2json_engine.py ./papers/ --api-key sk-xxx
"""

import argparse
import json
import os
import re
import sys
import time
from pathlib import Path
from typing import Optional

# ---------------------------------------------------------------------------
#  区块 1 — 文件装载读取器
# ---------------------------------------------------------------------------

SUPPORTED = {".pdf", ".docx"}


def extract_text_from_file(file_path: str) -> str:
    """按后缀分发引擎，返回拼接后的全量纯文本。"""
    ext = Path(file_path).suffix.lower()
    if ext not in SUPPORTED:
        raise ValueError(f"不支持的文件类型: {ext}（仅支持 {SUPPORTED}）")

    if ext == ".pdf":
        return _extract_pdf(file_path)
    return _extract_docx(file_path)


def _extract_pdf(path: str) -> str:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError("请安装 PyMuPDF: pip install PyMuPDF")

    doc = fitz.open(path)
    parts: list[str] = []
    for page in doc:
        text = page.get_text("text")
        if text.strip():
            parts.append(text.strip())
    doc.close()

    if not parts:
        # 回退：用 pdfplumber
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("PyMuPDF 提取为空，请尝试安装 pdfplumber")
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    parts.append(t.strip())
    return "\n\n".join(parts)


def _extract_docx(path: str) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("请安装 python-docx: pip install python-docx")

    doc = Document(path)
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    # 也抓表格里的文字
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


# ---------------------------------------------------------------------------
#  区块 2 — LLM 调用封装 (OpenAI 兼容协议)
# ---------------------------------------------------------------------------

SYSTEM_PROMPT = r"""
你是一个教辅题库拆解机器，请读取我提供的一堆杂乱文档。这些文档是一张或半张计算机科学科目的练习试卷（部分含有混入题中的干扰文或者无意义的标题页脚）。请剔除杂物，提取里面的具体题号信息，生成合法的标准 JSON Array 包裹。

每抽一道题目，结构字典定义如下(严格限制):
{
  "type": 0 或 1,
  "subject": "需推断这是考哪个大的课程大类(如数据结构、考研政治、408计网等)，并归类补充名称",
  "chapter": "请按前沿考点逻辑为本题编排并填入个章节分类词，以作知识库聚类",
  "content": "抽提出的真正原试卷的问题文字。请过滤没用的前导词如'题目42'，并将任何夹带混排出现的公式包裹进标准的一对 `$` 或块级两对 `$$`（用标准 LaTeX 语意修补还原输出！很重要！）",
  "options": "如果 type 是 0。此项务必是一个合法的 JSON 【数组序列化的字符串形态】。例形如：\"[\\\"A. XX\\\", \\\"B. YY\\\"]\"。绝对不可传 null，或者真正的 [ 结构过去。如果没有填入空型串: \"[]\"",
  "standard_answer": "结合文档内所随附或自己推理而来的单一或长正确文本",
  "analysis": "提供解答（或者依据自己的理智补充撰排清晰易懂、markdown排好分行甚至带着 $$代码快/树公式排版的分析），给后来考生看！"
}

必须生成标准的上述 Json list Array结构！不要夹带额外分析话语和首尾冗余段。
""".strip()


def call_llm_extractor(
    text_chunk: str,
    api_key: str,
    base_url: str = "https://api.openai.com/v1",
    model: str = "gpt-4o",
    timeout: int = 180,
    max_tokens: int = 16000,
) -> list[dict]:
    """将 text_chunk 发送给 LLM，返回解析后的题目列表。"""
    try:
        from openai import OpenAI
    except ImportError:
        raise ImportError("请安装 openai: pip install openai")

    client = OpenAI(api_key=api_key, base_url=base_url, timeout=timeout)

    resp = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": text_chunk},
        ],
        max_tokens=max_tokens,
        temperature=0.1,
        response_format={"type": "json_object"},
    )

    raw = resp.choices[0].message.content or ""
    return _parse_llm_output(raw)


def _parse_llm_output(raw: str) -> list[dict]:
    """容错解析 LLM 返回的 JSON 字符串。"""
    # 去掉 markdown 代码块包裹
    raw = raw.strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)

    try:
        data = json.loads(raw)
    except json.JSONDecodeError:
        # 尝试提取第一个 JSON 数组
        m = re.search(r"\[[\s\S]*\]", raw)
        if m:
            data = json.loads(m.group())
        else:
            return []

    # 兼容 {"questions": [...]} 包裹
    if isinstance(data, dict):
        for v in data.values():
            if isinstance(v, list):
                data = v
                break
        else:
            return []

    if not isinstance(data, list):
        return []

    return data


# ---------------------------------------------------------------------------
#  区块 3 — 清洗 & 入库标准化
# ---------------------------------------------------------------------------

def _sanitize_item(item: dict, default_subject: str, default_chapter: str) -> dict:
    """确保每条题目字段完整、类型正确。"""
    t = item.get("type", 0)
    if t not in (0, 1):
        t = 0

    options = item.get("options")
    if t == 0:
        if options is None or options == "":
            options = "[]"
        # 确保是字符串形态的 JSON 数组
        if isinstance(options, list):
            options = json.dumps(options, ensure_ascii=False)
        if isinstance(options, str):
            try:
                parsed = json.loads(options)
                if isinstance(parsed, list):
                    options = json.dumps(parsed, ensure_ascii=False)
                else:
                    options = "[]"
            except json.JSONDecodeError:
                options = "[]"
    else:
        options = "[]"

    return {
        "type": t,
        "subject": item.get("subject") or default_subject,
        "chapter": item.get("chapter") or default_chapter,
        "content": (item.get("content") or "").strip(),
        "options": options,
        "standard_answer": (item.get("standard_answer") or "").strip(),
        "analysis": (item.get("analysis") or "").strip(),
    }


def _batch_chunks(text: str, max_chars: int = 12000) -> list[str]:
    """将长文本切分为适合 LLM 上下文窗口的块，尽量在段落边界断开。"""
    if len(text) <= max_chars:
        return [text]

    paragraphs = text.split("\n\n")
    chunks: list[str] = []
    buf = ""
    for p in paragraphs:
        if len(buf) + len(p) + 2 > max_chars and buf:
            chunks.append(buf.strip())
            buf = p
        else:
            buf = buf + "\n\n" + p if buf else p
    if buf.strip():
        chunks.append(buf.strip())
    return chunks


# ---------------------------------------------------------------------------
#  主流程
# ---------------------------------------------------------------------------

def process_file(
    file_path: str,
    api_key: str,
    base_url: str,
    model: str,
    out_dir: str,
    default_subject: str,
    default_chapter: str,
    dry_run: bool,
) -> Optional[str]:
    print(f"\n{'=' * 60}")
    print(f"📄 处理: {file_path}")
    print(f"{'=' * 60}")

    # 1) 提取文本
    try:
        full_text = extract_text_from_file(file_path)
    except Exception as e:
        print(f"  ❌ 提取失败: {e}")
        return None

    if not full_text.strip():
        print("  ⚠️  提取到的文本为空，跳过")
        return None

    print(f"  📝 提取文本 {len(full_text)} 字符")

    # 2) 切块
    chunks = _batch_chunks(full_text)
    print(f"  🧩 切分为 {len(chunks)} 个块")

    if dry_run:
        print("  🔍 [dry-run] 不调用 LLM，仅展示文本预览:")
        for i, c in enumerate(chunks):
            print(f"  --- 块 {i+1} ({len(c)} chars) ---")
            print(c[:500])
            print("  ...")
        return None

    # 3) 逐块调用 LLM
    all_items: list[dict] = []
    for i, chunk in enumerate(chunks):
        print(f"  🤖 调用 LLM 处理块 {i+1}/{len(chunks)} …")
        try:
            items = call_llm_extractor(chunk, api_key, base_url, model)
        except Exception as e:
            print(f"  ⚠️  块 {i+1} 调用失败: {e}")
            continue

        if items:
            for item in items:
                all_items.append(_sanitize_item(item, default_subject, default_chapter))
            print(f"  ✅ 块 {i+1} 提取 {len(items)} 题")
        else:
            print(f"  ⚠️  块 {i+1} 未提取到题目")

        if i < len(chunks) - 1:
            time.sleep(1)  # 速率控制

    if not all_items:
        print("  ⚠️  未提取到任何题目")
        return None

    # 4) 最终包装为 Shiroha 兼容格式
    stem = Path(file_path).stem
    output = {
        "name": stem,
        "description": f"由 {stem} 自动抽取",
        "subject": default_subject,
        "chapter": default_chapter,
        "questions": all_items,
    }

    out_path = Path(out_dir) / f"{stem}_extracted.json"
    out_path.parent.mkdir(parents=True, exist_ok=True)
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(output, f, ensure_ascii=False, indent=2)

    print(f"  🎉 输出 {len(all_items)} 题 → {out_path}")
    return str(out_path)


def main():
    parser = argparse.ArgumentParser(
        description="Shiroha Quiz — PDF/DOCX → JSON 题库自动抽取引擎",
    )
    parser.add_argument(
        "input",
        nargs="+",
        help="输入文件或目录路径（支持 .pdf / .docx）",
    )
    parser.add_argument(
        "--api-key",
        default=os.environ.get("OPENAI_API_KEY", ""),
        help="API Key（默认读取环境变量 OPENAI_API_KEY）",
    )
    parser.add_argument(
        "--base-url",
        default=os.environ.get("OPENAI_BASE_URL", "https://api.openai.com/v1"),
        help="API Base URL（兼容 DeepSeek / GLM / 本地模型）",
    )
    parser.add_argument(
        "--model",
        default="gpt-4o",
        help="模型名称",
    )
    parser.add_argument(
        "--out-dir",
        default="./extracted/",
        help="输出目录",
    )
    parser.add_argument(
        "--subject",
        default="未分类科目",
        help="默认科目分类（JSON 文件中可覆盖）",
    )
    parser.add_argument(
        "--chapter",
        default="综合训练",
        help="默认章节分类",
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="仅提取文本预览，不调用 LLM",
    )
    args = parser.parse_args()

    if not args.api_key and not args.dry_run:
        print("❌ 请提供 --api-key 或设置环境变量 OPENAI_API_KEY")
        sys.exit(1)

    # 收集文件
    files: list[str] = []
    for path in args.input:
        p = Path(path)
        if p.is_dir():
            for ext in SUPPORTED:
                files.extend(str(x) for x in p.rglob(f"*{ext}"))
        elif p.suffix.lower() in SUPPORTED:
            files.append(str(p))
        else:
            print(f"⚠️  跳过不支持的文件: {path}")

    if not files:
        print("❌ 未找到可处理的文件")
        sys.exit(1)

    print(f"📂 共发现 {len(files)} 个文件")
    results = []
    for f in sorted(files):
        r = process_file(
            file_path=f,
            api_key=args.api_key,
            base_url=args.base_url,
            model=args.model,
            out_dir=args.out_dir,
            default_subject=args.subject,
            default_chapter=args.chapter,
            dry_run=args.dry_run,
        )
        if r:
            results.append(r)

    print(f"\n{'=' * 60}")
    print(f"🏁 完成！成功处理 {len(results)}/{len(files)} 个文件")
    for r in results:
        print(f"  → {r}")


if __name__ == "__main__":
    main()
